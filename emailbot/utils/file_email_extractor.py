"""Utilities to extract e-mail addresses from uploaded files."""

from __future__ import annotations

import csv
import html
import io
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Set, Tuple
from urllib.parse import unquote

from emailbot.ptb_context import get_current_chat_id
from emailbot.run_control import should_stop
from emailbot.ui.notify import (
    forget_timeout_hint_target,
    remember_timeout_hint_target,
)
from emailbot.utils.email_clean import clean_and_normalize_email
from emailbot.utils.logging_setup import get_logger
from emailbot.extraction_pdf import extract_emails_from_pdf as _extract_pdf_file

logger = get_logger(__name__)

SAFE_EMAIL_RE = re.compile(
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,63}",
    re.ASCII,
)
EMAIL_RE = SAFE_EMAIL_RE
CHUNK_BYTES = 128 * 1024
MAX_MATCHES_PER_FILE = 1000
MAX_BYTES = 25 * 1024 * 1024  # 25 MB per file
ALLOWED_IN_ZIP = {".txt", ".csv", ".tsv", ".pdf", ".docx", ".xlsx", ".htm", ".html"}


class ExtractError(Exception):
    """Raised when we cannot process a file."""


def _iter_chunks(text: str, chunk_bytes: int = CHUNK_BYTES) -> Iterable[str]:
    if not text:
        return
    encoded = text.encode("utf-8", errors="ignore")
    view = memoryview(encoded)
    for start in range(0, len(view), chunk_bytes):
        yield view[start : start + chunk_bytes].tobytes().decode("utf-8", errors="ignore")


def extract_emails_from_text(text: str) -> Set[str]:
    """Find potential e-mail addresses inside ``text`` using chunked scanning."""

    found: Set[str] = set()
    carry = ""
    for chunk in _iter_chunks(text):
        if carry:
            chunk = carry + chunk
            carry = ""
        if not chunk:
            continue
        parts = chunk.split()
        if chunk and not chunk[-1].isspace():
            carry = parts.pop() if parts else chunk
        for part in parts:
            if "@" not in part:
                continue
            for match in SAFE_EMAIL_RE.finditer(part):
                found.add(match.group(0))
                if len(found) >= MAX_MATCHES_PER_FILE:
                    return found
    if carry and "@" in carry:
        for match in SAFE_EMAIL_RE.finditer(carry):
            found.add(match.group(0))
            if len(found) >= MAX_MATCHES_PER_FILE:
                break
    return found


def _norm_and_dedupe(cands: List[str]) -> Tuple[List[str], Dict[str, int]]:
    ok: List[str] = []
    seen: set[str] = set()
    rejects: Dict[str, int] = {}
    for raw in cands:
        if should_stop():
            break
        email, reason = clean_and_normalize_email(raw)
        if email is None:
            key = str(reason) if reason else "unknown"
            rejects[key] = rejects.get(key, 0) + 1
            continue
        if email in seen:
            continue
        seen.add(email)
        ok.append(email)
    return ok, rejects


def _from_text(data: bytes) -> Tuple[List[str], Dict[str, int]]:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="ignore")
    emails = extract_emails_from_text(text)
    return _norm_and_dedupe(list(emails))


def _find_obfuscated_emails(text: str) -> List[str]:
    """
    Ищем реальный шаблон локаль + (at|@|собака) + домен (+ dot + …), но:
    - локаль должна содержать хотя бы ОДНУ букву и быть длиной ≥2;
    - каждый доменный лейбл должен содержать хотя бы ОДНУ букву,
      не начинаться/заканчиваться дефисом; TLD длиной 2–24.
    """

    # локальная часть: 1–64, обязательно ≥1 буква и длина ≥2
    local_part = r"(?P<local>(?=[A-Za-z0-9.+-]{2,64}$)(?=.*[A-Za-z])[A-Za-z0-9.+-]{2,64})"
    at_token = r"(?:\(|\[)?\s*(?:@|at|собака)\s*(?:\)|\])?"
    # доменный лейбл: 1–63, ≥1 буква, не начин/заканч с дефиса
    label = r"(?P<label>(?=[A-Za-z0-9-]{1,63}$)(?=.*[A-Za-z])[A-Za-z0-9](?:[A-Za-z0-9-]{0,61}[A-Za-z0-9])?)"
    dot_token = r"(?:\(|\[)?\s*(?:\.|dot|точка)\s*(?:\)|\])?"
    pattern = re.compile(
        rf"\b{local_part}\b\s*{at_token}\s*\b{label}\b(?:\s*{dot_token}\s*\b{label}\b)*",
        re.IGNORECASE | re.UNICODE,
    )

    emails: List[str] = []
    seen: set[str] = set()

    for match in pattern.finditer(text):
        if should_stop():
            break
        local = match.group("local")
        span_text = match.group(0)
        labels = re.findall(r"\b([A-Za-z0-9](?:[A-Za-z0-9-]{0,61}[A-Za-z0-9])?)\b", span_text)
        if not labels:
            continue
        parts = [p for p in labels if p.lower() != local.lower()]
        if not parts:
            continue
        domain = ".".join(parts)
        if "." not in domain:
            continue
        tld = domain.rsplit(".", 1)[-1]
        if not (2 <= len(tld) <= 24):
            continue
        labels_ok = True
        for lbl in domain.split("."):
            if should_stop():
                labels_ok = False
                break
            if not re.search(r"[A-Za-z]", lbl):
                labels_ok = False
                break
        if not labels_ok:
            continue
        email = f"{local}@{domain}"
        if email not in seen:
            seen.add(email)
            emails.append(email)
    return emails


def _from_html(data: bytes) -> Tuple[List[str], Dict[str, int]]:
    html_text = data.decode("utf-8", errors="ignore")
    mailto = re.findall(r"mailto:([^""' >]+)", html_text, flags=re.I)
    mailto = [unquote(m.split("?")[0]) for m in mailto]

    text = re.sub(r"<!--.*?-->", " ", html_text, flags=re.S)
    text = re.sub(r"<script\b[^>]*>.*?</script>", " ", text, flags=re.I | re.S)
    text = re.sub(r"<style\b[^>]*>.*?</style>", " ", text, flags=re.I | re.S)
    text = html.unescape(text)
    text = re.sub(r"<[^>]+>", " ", text)

    cands = set(EMAIL_RE.findall(text))
    cands.update(_find_obfuscated_emails(text))
    cands.update(mailto)
    return _norm_and_dedupe(list(cands))


def _from_csv(data: bytes) -> Tuple[List[str], Dict[str, int]]:
    buf = io.StringIO(data.decode("utf-8", errors="ignore"))
    sample = buf.read(4096)
    buf.seek(0)
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(buf, dialect)
    cands: List[str] = []
    for row in reader:
        if should_stop():
            break
        for cell in row:
            if should_stop():
                break
            if cell:
                cands.extend(EMAIL_RE.findall(cell))
    return _norm_and_dedupe(cands)


def _from_xlsx(data: bytes) -> Tuple[List[str], Dict[str, int], str | None]:
    try:
        import openpyxl  # type: ignore
    except Exception:
        return [], {"missing_dep_openpyxl": 1}, "Для .xlsx нужен пакет openpyxl"

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    cands: List[str] = []
    for ws in wb.worksheets:
        if should_stop():
            break
        for row in ws.iter_rows(values_only=True):
            if should_stop():
                break
            for val in row:
                if should_stop():
                    break
                if val is not None:
                    cands.extend(EMAIL_RE.findall(str(val)))
    ok, rej = _norm_and_dedupe(cands)
    return ok, rej, None


def _from_docx(data: bytes) -> Tuple[List[str], Dict[str, int], str | None]:
    try:
        import docx  # type: ignore
    except Exception:
        return [], {"missing_dep_python_docx": 1}, "Для .docx нужен пакет python-docx"

    document = docx.Document(io.BytesIO(data))
    texts = [p.text for p in document.paragraphs]
    for table in document.tables:
        if should_stop():
            break
        for row in table.rows:
            if should_stop():
                break
            for cell in row.cells:
                if should_stop():
                    break
                texts.append(cell.text)
    cands: List[str] = []
    for text in texts:
        if should_stop():
            break
        cands.extend(EMAIL_RE.findall(text))
    ok, rej = _norm_and_dedupe(cands)
    return ok, rej, None


def _from_pdf(data: bytes) -> Tuple[List[str], Dict[str, int], str | None]:
    if not data:
        return [], {}, None

    tmp: tempfile.NamedTemporaryFile | None = None
    tmp_path: Path | None = None
    emails: set[str] | None = None
    timeout_key: Path | None = None
    try:
        tmp = tempfile.NamedTemporaryFile(prefix="ebot_pdf_", suffix=".pdf", delete=False)
        tmp.write(data)
        tmp.flush()
        tmp_path = Path(tmp.name)
        logger.info("PDF bridge: saved stream -> %s", tmp_path.name)
        try:
            chat_id = get_current_chat_id()
        except Exception:
            chat_id = None
        if chat_id is not None:
            remember_timeout_hint_target(tmp_path, chat_id)
            timeout_key = tmp_path
        emails = _extract_pdf_file(tmp_path)
    except Exception:
        logger.exception("PDF bridge failed")
    finally:
        if timeout_key is not None:
            try:
                forget_timeout_hint_target(timeout_key)
            except Exception:
                pass
        if tmp is not None:
            try:
                tmp.close()
            except Exception:
                pass
        if tmp_path is not None:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass

    if emails is not None:
        ok = sorted(emails)
        return ok, {}, None

    try:
        from pdfminer.high_level import extract_text
    except Exception:
        return [], {"missing_dep_pdfminer": 1}, "Для .pdf нужен пакет pdfminer.six"

    if should_stop():
        return [], {}, None
    text = extract_text(io.BytesIO(data), maxpages=200) or ""
    ok, rej = _norm_and_dedupe(EMAIL_RE.findall(text))
    return ok, rej, None


def _safe_zip_name(name: str) -> bool:
    return not (name.startswith("/") or ".." in name.replace("\\", "/"))


def _from_zip(data: bytes) -> Tuple[List[str], Dict[str, int], List[str]]:
    ok_all: List[str] = []
    rejects: Dict[str, int] = {}
    errors: List[str] = []

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for info in zf.infolist():
            if should_stop():
                break
            filename = info.filename
            if info.is_dir() or not _safe_zip_name(filename):
                if not _safe_zip_name(filename):
                    errors.append(f"{filename}: пропуск (путь)")
                continue
            if info.file_size > MAX_BYTES:
                errors.append(f"{filename}: пропуск (> {MAX_BYTES // (1024 * 1024)}MB)")
                continue
            ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            if ext not in ALLOWED_IN_ZIP:
                continue
            with zf.open(info) as fh:
                content = fh.read()
            if should_stop():
                break
            ok, rej, warn = extract_emails_from_bytes(content, filename)
            ok_all.extend(ok)
            for key, val in rej.items():
                rejects[key] = rejects.get(key, 0) + val
            if warn:
                errors.append(f"{filename}: {warn}")

    ok_all = list(dict.fromkeys(ok_all))
    return ok_all, rejects, errors


def extract_from_plain_text(text: str) -> Set[str]:
    """Public helper for quickly extracting e-mail addresses from ``text``."""

    return extract_emails_from_text(text)


def extract_emails_from_bytes(data: bytes, filename: str) -> Tuple[List[str], Dict[str, int], str | None]:
    """Extract addresses from ``data`` according to the file name."""

    if len(data) > MAX_BYTES:
        raise ExtractError("file_too_large")

    name = filename.lower()
    if name.endswith((".csv", ".tsv")):
        ok, rejects = _from_csv(data)
        return ok, rejects, None
    if name.endswith(".txt"):
        ok, rejects = _from_text(data)
        return ok, rejects, None
    if name.endswith((".html", ".htm")):
        ok, rejects = _from_html(data)
        return ok, rejects, None
    if name.endswith(".xlsx"):
        return _from_xlsx(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".zip"):
        ok_all, rejects, errors = _from_zip(data)
        warn = "; ".join(errors) if errors else None
        return ok_all, rejects, warn

    ok, rejects = _from_text(data)
    return ok, rejects, None
