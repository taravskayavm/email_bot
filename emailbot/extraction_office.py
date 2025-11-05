from pathlib import Path
from typing import Set


def extract_emails_from_docx(path: Path) -> Set[str]:
    try:
        import docx  # python-docx
        from emailbot.parsing.extract_from_text import emails_from_text
    except Exception:
        return set()
    try:
        doc = docx.Document(str(path))
        buf: list[str] = []
        for p in doc.paragraphs:
            buf.append(p.text or "")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    buf.append(cell.text or "")
        return emails_from_text("\n".join(buf))
    except Exception:
        return set()


def extract_emails_from_xlsx(path: Path) -> Set[str]:
    try:
        from openpyxl import load_workbook
        from emailbot.parsing.extract_from_text import emails_from_text
    except Exception:
        return set()
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        found: Set[str] = set()
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for v in row:
                    if v:
                        found |= emails_from_text(str(v))
        return found
    except Exception:
        return set()


def extract_emails_from_xls(path: Path) -> Set[str]:
    # поддержка старого .xls, если установлен xlrd<=1.2 или pyxlsb не годится
    try:
        import xlrd  # требуется старая версия для .xls
        from emailbot.parsing.extract_from_text import emails_from_text
    except Exception:
        return set()
    try:
        book = xlrd.open_workbook(str(path))
        found: Set[str] = set()
        for sh in book.sheets():
            for r in range(sh.nrows):
                for c in range(sh.ncols):
                    v = sh.cell_value(r, c)
                    if v:
                        found |= emails_from_text(str(v))
        return found
    except Exception:
        return set()
