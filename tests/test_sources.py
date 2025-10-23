import asyncio
import base64
import http.server
import os
import socket
import threading
import zipfile
from pathlib import Path

import pytest

pytest.importorskip("emailbot.bot_handlers")

from emailbot import extraction


def _create_pdf(path: Path, text: str) -> None:
    import fitz  # type: ignore

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(path)


def _create_pdf_with_mailto(path: Path, email: str) -> None:
    import fitz  # type: ignore

    doc = fitz.open()
    page = doc.new_page()
    # текст без явного email, чтобы поймать mailto-аннотацию
    page.insert_text((72, 72), "Contact us")
    rect = fitz.Rect(72, 72, 200, 90)
    page.insert_link({
        "kind": fitz.LINK_URI,
        "from": rect,
        "uri": f"mailto:{email}?subject=Hello",
    })
    doc.save(path)


def _create_docx(path: Path, text: str) -> None:
    import docx  # type: ignore

    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(path)


def _create_xlsx(path: Path, text: str) -> None:
    import openpyxl  # type: ignore

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([text])
    wb.save(path)


def _run_server(pages: dict[str, str]):
    class Handler(http.server.BaseHTTPRequestHandler):
        def do_GET(self):  # noqa: N802
            body = pages.get(self.path, "").encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def log_message(self, *a):  # noqa: D401, override
            return

    sock = socket.socket()
    sock.bind(("localhost", 0))
    port = sock.getsockname()[1]
    sock.close()
    server = http.server.ThreadingHTTPServer(("localhost", port), Handler)
    thread = threading.Thread(target=server.serve_forever)
    thread.daemon = True
    thread.start()
    return server, port


def test_extract_emails_from_zip(tmp_path: Path):
    pdf = tmp_path / "a.pdf"
    docx = tmp_path / "b.docx"
    xlsx = tmp_path / "c.xlsx"
    csv = tmp_path / "d.csv"
    _create_pdf(pdf, "pdf@example.com")
    _create_docx(docx, "docx@example.com")
    _create_xlsx(xlsx, "xlsx@example.com")
    csv.write_text("email\ncsv@example.com\n", encoding="utf-8")

    zip_path = tmp_path / "test.zip"
    with zipfile.ZipFile(zip_path, "w") as z:
        for f in [pdf, docx, xlsx, csv]:
            z.write(f, f.name)

    hits, stats = extraction.extract_emails_from_zip(str(zip_path))
    emails = [h.email for h in hits]
    assert set(emails) == {
        "pdf@example.com",
        "docx@example.com",
        "xlsx@example.com",
        "csv@example.com",
    }
    assert stats.get("pdf") == 1
    assert stats.get("docx") == 1
    assert stats.get("xlsx") == 1
    assert stats.get("csv") == 1


def test_extract_from_url(tmp_path: Path):
    # cfemail encoding for test@example.com
    key = 0x12
    cfemail = f"{key:02x}" + "".join(
        f"{b ^ key:02x}" for b in "hidden@example.com".encode("utf-8")
    )
    pages = {
        "/": (
            '<a href="mailto:hello@site.ru">x</a> '
            "user@domain.ru "
            "user [at] example [dot] com "
            "user (at) test (dot) ru "
            f'<a data-cfemail="{cfemail}"></a>'
            '<a href="/about/contacts">next</a>'
        ),
        "/about/contacts": "contact@site.ru",
        "/empty": "<html><body>No mail</body></html>",
    }
    server, port = _run_server(pages)
    try:
        hits, stats = extraction.extract_from_url(f"http://localhost:{port}/")
        emails = [h.email for h in hits]
        assert "hello@site.ru" in emails
        assert "user@domain.ru" in emails
        assert "user@example.com" in emails
        assert "user@test.ru" in emails
        assert "hidden@example.com" in emails
        assert stats["cfemail_decoded"] == 1
        assert stats["obfuscated_hits"] >= 2
        assert stats["urls_scanned"] >= 1

        hits_empty, stats_empty = extraction.extract_from_url(
            f"http://localhost:{port}/empty"
        )
        assert hits_empty == []
        assert stats_empty["urls_scanned"] == 1
    finally:
        server.shutdown()


def test_extract_from_documents(tmp_path: Path):
    pdf = tmp_path / "f.pdf"
    docx = tmp_path / "f.docx"
    xlsx = tmp_path / "f.xlsx"
    txt = tmp_path / "f.txt"
    _create_pdf(pdf, "a@pdf.com")
    _create_docx(docx, "b@docx.com")
    _create_xlsx(xlsx, "c@xlsx.com")
    txt.write_text("d@text.com", encoding="utf-8")

    hits_pdf, _ = extraction.extract_from_pdf(str(pdf))
    hits_docx, _ = extraction.extract_from_docx(str(docx))
    hits_xlsx, _ = extraction.extract_from_xlsx(str(xlsx))
    hits_txt, _ = extraction.extract_from_csv_or_text(str(txt))
    assert [h.email for h in hits_pdf] == ["a@pdf.com"]
    assert [h.email for h in hits_docx] == ["b@docx.com"]
    assert [h.email for h in hits_xlsx] == ["c@xlsx.com"]
    assert [h.email for h in hits_txt] == ["d@text.com"]


def test_extract_mailto_from_pdf(tmp_path: Path):
    pdf = tmp_path / "mailto.pdf"
    _create_pdf_with_mailto(pdf, "hello@example.com")

    hits, _ = extraction.extract_from_pdf(str(pdf))
    emails = {h.email for h in hits}

    assert "hello@example.com" in emails


def test_xlsx_no_handle_leak(tmp_path: Path):
    path = tmp_path / "leak.xlsx"
    _create_xlsx(path, "leak@example.com")
    extraction.extract_from_xlsx(str(path))
    os.remove(path)


@pytest.mark.asyncio
async def test_zip_handler_signature(monkeypatch, tmp_path: Path):
    from tests.test_bot_handlers import DummyContext, DummyDocument, DummyUpdate
    import emailbot.bot_handlers as bh

    update = DummyUpdate(document=DummyDocument())
    update.message.document.file_name = "data.zip"
    ctx = DummyContext()

    monkeypatch.setattr(bh, "DOWNLOAD_DIR", tmp_path)

    called = {}

    async def fake(path):
        called["arg"] = path
        return set(), [], set(), {}

    monkeypatch.setattr(bh, "extract_emails_from_zip", fake)

    await bh.handle_document(update, ctx)
    assert called["arg"].endswith("data.zip")

