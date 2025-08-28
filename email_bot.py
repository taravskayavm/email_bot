# -*- coding: utf-8 -*-
# email_bot.py — версия с устранением «усечённых» адресов и улучшенной склейкой

import os
import re
import csv
import time
import ssl
import asyncio
import zipfile
import imaplib
import smtplib
import email
import random
import threading
import concurrent.futures
from datetime import datetime, timedelta
from email.message import EmailMessage
from email.utils import formataddr
import html as htmllib
from pathlib import Path
from typing import Tuple, Set, List, Dict

import aiohttp
import fitz                       # PyMuPDF
import pandas as pd
from docx import Document
from telegram import (
    Update, ReplyKeyboardMarkup, InlineKeyboardButton, InlineKeyboardMarkup
)
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler,
    ContextTypes, filters
)

# ---------------- Пути/Конфиг ----------------
SCRIPT_DIR = Path(__file__).resolve().parent

DOWNLOAD_DIR = str(SCRIPT_DIR / "downloads")
LOG_FILE = str(SCRIPT_DIR / "sent_log.csv")
BLOCKED_FILE = str(SCRIPT_DIR / "blocked_emails.txt")
MAX_EMAILS_PER_DAY = 200

TEMPLATES_DIR = str(SCRIPT_DIR / "templates")
TEMPLATE_MAP = {
    "спорт":    os.path.join(TEMPLATES_DIR, "sport.htm"),
    "туризм":   os.path.join(TEMPLATES_DIR, "tourism.htm"),
    "медицина": os.path.join(TEMPLATES_DIR, "medicine.htm"),
}
SIGNATURE_HTML = (
    '<div style="margin-top:20px;font-size:12px;color:#666">'
    '—<br>Если вы больше не хотите получать письма — ответьте на это письмо словом <b>Unsubscribe</b>.'
    "</div>"
)

PRIVACY_NOTICE_HTML = (
    '<div style="margin-top:16px;font:12px/1.4 -apple-system,Segoe UI,Roboto,Arial,sans-serif;color:#666">'
    '<div style="border-top:1px solid #e5e5e5;margin:12px 0 8px"></div>'
    '<b>Почему вы получили это письмо?</b>'
    '<div>Мы пишем по профессиональному адресу, опубликованному в открытых источниках '
    '(официальные сайты/профили публикаций), с предложением профильного сотрудничества.</div>'
    '<div style="margin-top:6px"><b>Правовое основание:</b> legitimate interests (ст. 6(1)(f) GDPR / '
    'профильные интересы в РФ). <b>Цели:</b> экспертное/издательское сотрудничество. '
    '<b>Источник:</b> публичные страницы организации/автора.</div>'
    '<div style="margin-top:6px"><b>Ваши права:</b> вы можете возразить против подобных писем и/или отписаться — '
    'ответьте <b>Unsubscribe</b> на это письмо; мы добавим адрес в список исключений. '
    'Срок хранения контакта — не более необходимого для коммуникации, записи об отписке — дольше, '
    'чтобы не писать повторно.</div>'
    '<div style="margin-top:6px">Политику конфиденциальности и контакты для запросов можно получить по запросу.</div>'
    '</div>'
)

TECH_PATTERNS = ["noreply", "no-reply", "do-not-reply", "donotreply",
                 "postmaster", "mailer-daemon", "abuse", "support", "admin", "info@"]

ALLOWED_TLDS = {"ru", "com"}  # только эти домены отправляем

PREVIEW_ALLOWED = 10
PREVIEW_NUMERIC = 6
PREVIEW_FOREIGN = 6

FORCE_SEND_CHAT_IDS: set[int] = set()
session_data: dict[int, dict] = {}

_recent_cache = {"at": None, "set": set(), "ttl": 600}

TOKEN = ""
EMAIL_ADDRESS = ""
EMAIL_PASSWORD = ""


# ---------------- Загрузка .env ----------------
def load_env():
    try:
        from dotenv import load_dotenv as _load
        _load(dotenv_path=SCRIPT_DIR / ".env")
        _load()
    except Exception:
        pass

    def read_env_file(p: Path):
        if not p.exists():
            return
        try:
            with p.open(encoding="utf-8") as f:
                for line in f:
                    s = line.strip()
                    if not s or s.startswith("#") or "=" not in s:
                        continue
                    k, v = s.split("=", 1)
                    k = k.strip()
                    v = v.strip().strip('\'"')
                    os.environ.setdefault(k, v)
        except Exception:
            pass

    read_env_file(SCRIPT_DIR / ".env")
    read_env_file(Path.cwd() / ".env")


# ---------------- Утилиты ----------------
def normalize_email(s: str) -> str:
    return (s or "").strip().lower()


def is_allowed_tld(email_addr: str) -> bool:
    e = normalize_email(email_addr)
    return bool(re.search(r'@[A-Za-z0-9.-]+\.(?:ru|com)$', e))


def strip_html(html: str) -> str:
    if not html:
        return ""
    s = html
    s = re.sub(r'(?is)<(script|style).*?>.*?</\1>', '', s)
    s = re.sub(r'(?i)<br\s*/?>', '\n', s)
    s = re.sub(r'(?i)</p\s*>', '\n', s)
    s = re.sub(r'(?i)</div\s*>', '\n', s)
    s = re.sub(r'(?is)<[^>]+>', ' ', s)
    s = htmllib.unescape(s)
    s = re.sub(r'[ \t]+', ' ', s)
    s = re.sub(r'\n\s*\n+', '\n', s)
    return s.strip()


def log_error(msg: str):
    try:
        with open(SCRIPT_DIR / "bot_errors.log", "a", encoding="utf-8") as f:
            f.write(f"{datetime.now().isoformat()} {msg}\n")
    except Exception:
        pass


def sample_preview(items, k: int) -> list[str]:
    lst = list(dict.fromkeys(items))
    if len(lst) <= k:
        return lst
    return random.sample(lst, k)


# ---------------- Предобработка текста ----------------
def _preclean_text_for_emails(text: str) -> str:
    if not text:
        return ""
    s = text

    # убираем невидимые
    s = s.replace("\u00ad", "").replace("\u2011", "").replace("\u200b", "")
    s = s.replace("\xa0", " ")

    # защита от прилипания односимвольных маркеров перед email
    s = re.sub(
        r'(?im)\b([A-Za-z0-9])\s*[\-\)\]\u2010\u2011\u2012\u2013\u2014]\s*\n\s*(?=[A-Za-z][A-Za-z0-9._%+-]*@)',
        '',
        s
    )

    # де-гипенизация переносов (g-\nmail → gmail)
    s = re.sub(
        r'([A-Za-z0-9._%+\-])[\-\u2010\u2011\u2012\u2013\u2014]\s*\n\s*([A-Za-z0-9._%+\-])',
        r'\1\2',
        s
    )

    # склейка без дефиса — буква на новой строке
    s = re.sub(r'([A-Za-z]{3,})\s*\n\s*([A-Za-z][A-Za-z0-9._%+\-]*)@', r'\1\2@', s)
    s = re.sub(r'([A-Za-z]{2,})([._])\s*\n\s*([A-Za-z][A-Za-z0-9._%+\-]*)@', r'\1\2\3@', s)

    # новый кейс: слово на строке + ЧИСЛА на следующей + (возможные) пробелы перед '@'
    s = re.sub(r'([A-Za-z]{2,})\s*\n\s*([0-9]{1,6})\s*@', r'\1\2@', s)

    # \r/\n -> пробел
    s = re.sub(r'[\r\n]+', ' ', s)

    # убрать пробелы вокруг '@' и точки
    s = re.sub(r'\s*@\s*', '@', s)
    s = re.sub(r'(@[A-Za-z0-9.-]+)\s*\.\s*([A-Za-z]{2,10})\b', r'\1.\2', s)

    # '. c o m' / '. r u'
    s = re.sub(r'\.\s*c\s*o\s*m\b', '.com', flags=re.I, string=s)
    s = re.sub(r'\.\s*r\s*u\b', '.ru',  flags=re.I, string=s)

    # '@gmail.co' → '.com' (и др. провайдеры)
    prov = r"(gmail|yahoo|hotmail|outlook|protonmail|icloud|aol|live|msn|mail|yandex|rambler|bk|list|inbox|ya)"
    s = re.sub(rf'(@{prov}\.co)(?=[^\w]|$)', r'@\1.com', s, flags=re.I)
    s = re.sub(rf'(@{prov}\.co)\s*m\b', r'@\1.com', s, flags=re.I)

    # разделим «слипшийся хвост» после .ru/.com
    s = re.sub(r'(\.(?:ru|com))(?=[A-Za-z0-9])', r'\1 ', s)

    return s


# ---------------- Извлечение email ----------------
def extract_emails_loose(text: str) -> List[str]:
    if not text:
        return []
    s = _preclean_text_for_emails(text)
    rx = re.compile(r'([A-Za-z0-9][A-Za-z0-9._%+-]*@[A-Za-z0-9.-]+\.[A-Za-z]{2,})')
    return [normalize_email(x) for x in rx.findall(s)]


def collapse_footnote_variants(emails: set[str]) -> set[str]:
    if not emails:
        return set()
    base = {re.sub(r'^\.+', '', normalize_email(e)) for e in emails}
    by_suffix: dict[str, set[str]] = {}
    prefix_of: dict[str, str] = {}
    for e in list(base):
        m_num = re.match(r'^(\d{1,2})([A-Za-z][A-Za-z0-9._%+-]*@.+)$', e, flags=re.I)
        if m_num:
            by_suffix.setdefault(m_num.group(2), set()).add(e)
            prefix_of[e] = m_num.group(1); continue
        m_chr = re.match(r'^([A-Za-z])([A-Za-z][A-Za-z0-9._%+-]*@.+)$', e, flags=re.I)
        if m_chr:
            by_suffix.setdefault(m_chr.group(2), set()).add(e)
            prefix_of[e] = m_chr.group(1); continue
    keep = set(base)
    for suffix, variants in by_suffix.items():
        clean_present = suffix in keep
        distinct_pfx = set(prefix_of[v] for v in variants if v in prefix_of)
        if clean_present or len(distinct_pfx) >= 2:
            keep.difference_update(variants)
            keep.add(suffix)
    return keep


def extract_clean_emails_from_text(text: str) -> Set[str]:
    if not text:
        return set()
    text = _preclean_text_for_emails(text)
    base_re = re.compile(r'([A-Za-z0-9][A-Za-z0-9._%+-]*@[A-Za-z0-9.-]+\.(?:ru|com))(?=[^\w]|$)')
    raw = set(base_re.findall(text))
    if not raw:
        return set()
    result: Set[str] = {re.sub(r'^\.+', '', e) for e in raw}
    result = collapse_footnote_variants(result)
    result = {e for e in result if is_allowed_tld(e)}
    return result


def is_numeric_localpart(email_addr: str) -> bool:
    e = normalize_email(email_addr)
    return "@" in e and e.split("@", 1)[0].isdigit()


# ---------- Поиск/устранение усечённых адресов (33@ → vilena33@) ----------
def detect_numeric_truncations(candidates: Set[str]) -> List[tuple[str, str]]:
    """
    Ищем пары (bad, good), где bad = '<digits>@domain', а good = '<letters+digits>@same_domain'.
    Берём только случаи с ОДНОЗНАЧНЫМ соответствием.
    """
    by_key: Dict[tuple[str, str], Set[str]] = {}
    for e in candidates:
        loc, dom = e.split("@", 1)
        m = re.match(r'^([a-z]+)(\d{1,6})$', loc)
        if m:
            key = (m.group(2), dom)
            by_key.setdefault(key, set()).add(e)

    pairs: List[tuple[str, str]] = []
    for e in list(candidates):
        loc, dom = e.split("@", 1)
        if loc.isdigit():
            key = (loc, dom)
            fulls = by_key.get(key, set())
            if len(fulls) == 1:               # однозначное совпадение
                good = next(iter(fulls))
                pairs.append((e, good))
    return pairs


def apply_numeric_truncation_removal(allowed_set: Set[str]) -> Tuple[Set[str], List[tuple[str, str]]]:
    """
    Удаляет 'цифровые' логины, если найден ровно один длинный вариант.
    Возвращает (очищенный_набор, список_пар bad→good).
    """
    pairs = detect_numeric_truncations(allowed_set)
    if not pairs:
        return allowed_set, []
    cleaned = set(allowed_set)
    for bad, _ in pairs:
        cleaned.discard(bad)
    return cleaned, pairs


def _read_template_file(path: str) -> str:
    if not os.path.exists(path):
        alt = os.path.splitext(path)[0] + ".html"
        if os.path.exists(alt):
            path = alt
    with open(path, encoding="utf-8") as f:
        return f.read()


# --- извлечение из разных форматов (возвращаем два множества: allowed, loose_all) ---
def _extract_from_pdf(path: str) -> Tuple[Set[str], Set[str]]:
    doc = fitz.open(path)
    texts = []
    for page in doc:
        page_text = page.get_text() or ""
        texts.append(page_text)
    doc.close()
    joined = " ".join(texts)
    loose = set(extract_emails_loose(joined))
    allowed = set(extract_clean_emails_from_text(joined))
    return allowed, loose


def _extract_from_docx(path: str) -> Tuple[Set[str], Set[str]]:
    doc = Document(path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    loose = set(extract_emails_loose(full_text))
    allowed = set(extract_clean_emails_from_text(full_text))
    return allowed, loose


def _extract_from_excel(path: str) -> Tuple[Set[str], Set[str]]:
    emails_allowed, emails_loose = set(), set()
    try:
        df = pd.read_excel(path, dtype=str)
        for col in df.columns:
            for val in df[col].dropna():
                s = str(val)
                emails_allowed.update(extract_clean_emails_from_text(s))
                emails_loose.update(extract_emails_loose(s))
    except Exception as e:
        log_error(f"extract_from_excel: {path}: {e}")
    return emails_allowed, emails_loose


def _extract_from_csv(path: str) -> Tuple[Set[str], Set[str]]:
    emails_allowed, emails_loose = set(), set()
    try:
        df = pd.read_csv(path, header=None, dtype=str)
        for col in df.columns:
            for val in df[col].dropna():
                s = str(val)
                emails_allowed.update(extract_clean_emails_from_text(s))
                emails_loose.update(extract_emails_loose(s))
    except Exception as e:
        log_error(f"extract_from_csv: {path}: {e}")
    return emails_allowed, emails_loose


def extract_from_uploaded_file(path: str) -> Tuple[Set[str], Set[str]]:
    p = path.lower()
    if p.endswith(".pdf"):
        return _extract_from_pdf(path)
    if p.endswith(".xlsx"):
        return _extract_from_excel(path)
    if p.endswith(".csv"):
        return _extract_from_csv(path)
    if p.endswith(".docx"):
        return _extract_from_docx(path)
    return set(), set()


async def async_extract_emails_from_url(url: str, session, chat_id: int | None = None):
    try:
        async with session.get(url, timeout=20) as resp:
            html_text = await resp.text()
            allowed = extract_clean_emails_from_text(html_text)
            loose = set(extract_emails_loose(html_text))
            foreign = {e for e in loose if not is_allowed_tld(e)}
            # из html получаем кандидаты «буква/цифры перед @» в блок исправлений
            repairs = find_prefix_repairs(html_text)
            return (url, list(allowed), list(foreign), repairs)
    except Exception as e:
        log_error(f"async_extract_emails_from_url: {url}: {e}")
        return (url, [], [], [])


# ---------- Repairs (буква на строке/слово+цифры перед '@') ----------
def _remove_invisibles_keep_newlines(text: str) -> str:
    if not text:
        return ""
    s = text
    s = s.replace("\u00ad", "").replace("\u2011", "").replace("\u200b", "")
    s = s.replace("\xa0", " ")
    return s


def find_prefix_repairs(raw_text: str) -> List[tuple[str, str]]:
    """
    Возвращаем пары (bad_email, fixed_email) для .ru/.com:
      A) 'm\\norgachov-ilya@yandex.ru' → orgachov… → morgachov…
      B) 'Vilena\\n33 @mail.ru'       → 33@…       → vilena33@…
    """
    if not raw_text:
        return []
    s = _remove_invisibles_keep_newlines(raw_text)
    pairs, seen = [], set()

    pat_a = re.compile(r'(?im)\b([a-z])\s*\n\s*([a-z][a-z0-9._%+\-]{2,})@([a-z0-9.-]+\.(?:ru|com))')
    for m in pat_a.finditer(s):
        left, rest, dom = m.group(1).lower(), m.group(2).lower(), m.group(3).lower()
        bad, good = f"{rest}@{dom}", f"{left}{rest}@{dom}"
        if (bad, good) not in seen:
            seen.add((bad, good)); pairs.append((bad, good))

    pat_b = re.compile(r'(?im)\b([a-z]{2,})\s*\n\s*([0-9]{1,6})\s*@([a-z0-9.-]+\.(?:ru|com))')
    for m in pat_b.finditer(s):
        word, digits, dom = m.group(1).lower(), m.group(2), m.group(3).lower()
        bad, good = f"{digits}@{dom}", f"{word}{digits}@{dom}"
        if (bad, good) not in seen:
            seen.add((bad, good)); pairs.append((bad, good))

    return pairs


def collect_repairs_from_files(file_paths: List[str]) -> List[tuple[str, str]]:
    repairs: List[tuple[str, str]] = []
    for path in file_paths:
        p = path.lower()
        try:
            if p.endswith(".pdf"):
                doc = fitz.open(path)
                try:
                    raw = "\n".join((pg.get_text() or "") for pg in doc)
                finally:
                    doc.close()
                repairs.extend(find_prefix_repairs(raw))
            elif p.endswith(".docx"):
                doc = Document(path)
                raw = "\n".join(para.text for para in doc.paragraphs)
                repairs.extend(find_prefix_repairs(raw))
        except Exception as e:
            log_error(f"collect_repairs_from_files: {path}: {e}")
    uniq = list(dict.fromkeys(repairs))
    return uniq


def extract_emails_multithreaded(file_paths: List[str]) -> Tuple[Set[str], Set[str]]:
    allowed_all, loose_all = set(), set()
    def process(file):
        try:
            return extract_from_uploaded_file(file)
        except Exception as ex:
            log_error(f"extract_emails_multithreaded:{file}: {ex}")
            return set(), set()
    with concurrent.futures.ThreadPoolExecutor() as executor:
        for allowed, loose in executor.map(process, file_paths):
            allowed_all.update(allowed); loose_all.update(loose)
    return allowed_all, loose_all


async def extract_emails_from_zip(zip_path: str, progress_msg, download_dir: str) -> Tuple[Set[str], List[str], Set[str]]:
    all_allowed: Set[str] = set()
    all_loose: Set[str] = set()
    extracted_files: List[str] = []
    with zipfile.ZipFile(zip_path, 'r') as z:
        file_list = [f for f in z.namelist() if f.lower().endswith(('.pdf', '.xlsx', '.csv', '.docx'))]
        total_files = len(file_list)
        if total_files == 0:
            if progress_msg:
                await progress_msg.edit_text("❌ В архиве не найдено поддерживаемых файлов.")
            return all_allowed, extracted_files, all_loose
        if progress_msg:
            await progress_msg.edit_text(f"В архиве {total_files} файлов. Начинаем обработку...")
        for idx, inner_file in enumerate(file_list, 1):
            extracted_path = os.path.join(download_dir, inner_file)
            os.makedirs(os.path.dirname(extracted_path), exist_ok=True)
            z.extract(inner_file, download_dir)
            extracted_files.append(extracted_path)
            allowed, loose = extract_from_uploaded_file(extracted_path)
            all_allowed.update(allowed); all_loose.update(loose)
            if progress_msg:
                await progress_msg.edit_text(f"🔄 Обработка файла {idx}/{total_files}:\n{inner_file}")
    return all_allowed, extracted_files, all_loose


# ---------------- SMTP/Отправка ----------------
def send_raw_smtp_with_retry(raw_message: str, recipient: str, max_tries=3):
    last_exc = None
    for _ in range(max_tries):
        try:
            with smtplib.SMTP_SSL("smtp.mail.ru", 465, context=ssl.create_default_context()) as server:
                server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
                server.sendmail(EMAIL_ADDRESS, recipient, raw_message)
                return
        except Exception as e:
            last_exc = e; time.sleep(2)
    raise last_exc


def save_to_sent_folder(raw_message: str):
    try:
        imap = imaplib.IMAP4_SSL("imap.mail.ru")
        imap.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        sent_folder = detect_sent_folder(imap)
        imap.append(f'"{sent_folder}"', '\\Seen', imaplib.Time2Internaldate(time.time()), raw_message.encode('utf-8'))
        imap.logout()
    except Exception as e:
        log_error(f"save_to_sent_folder: {e}")


def build_message(to_addr: str, html_path: str, subject: str, extra_html: str | None = None) -> EmailMessage:
    html_body = _read_template_file(html_path)
    html_body = html_body.replace("</body>", f"{SIGNATURE_HTML}</body>")
    if extra_html:
        html_body = html_body.replace("</body>", f"{extra_html}</body>")
    text_body = strip_html(html_body)
    msg = EmailMessage()
    msg["From"] = formataddr(("Редакция литературы по медицине, спорту и туризму", EMAIL_ADDRESS))
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg["Reply-To"] = EMAIL_ADDRESS
    msg["List-Unsubscribe"] = f"<mailto:{EMAIL_ADDRESS}?subject=unsubscribe>"
    msg["List-Unsubscribe-Post"] = "List-Unsubscribe=One-Click"
    msg.set_content(text_body)
    msg.add_alternative(html_body, subtype="html")
    logo_path = SCRIPT_DIR / "Logo.png"
    if logo_path.exists():
        try:
            with logo_path.open("rb") as img:
                img_bytes = img.read()
            msg.get_payload()[-1].add_related(img_bytes, maintype="image", subtype="png", cid="<logo>")
        except Exception as e:
            log_error(f"attach_logo: {e}")
    return msg


def _is_first_contact(recipient: str) -> bool:
    recent = get_recent_6m_union()
    return normalize_email(recipient) not in recent


def send_email(recipient: str, html_path: str, subject: str = "Издательство Лань приглашает к сотрудничеству", notify_func=None):
    try:
        extra_html = PRIVACY_NOTICE_HTML if _is_first_contact(recipient) else None
        msg = build_message(recipient, html_path, subject, extra_html=extra_html)
        raw = msg.as_string()
        send_raw_smtp_with_retry(raw, recipient, max_tries=3)
        save_to_sent_folder(raw)
    except Exception as e:
        log_error(f"send_email: {recipient}: {e}")
        if notify_func:
            notify_func(f"❌ Ошибка при отправке на {recipient}: {e}")
        raise


async def async_send_email(recipient: str, html_path: str):
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, send_email, recipient, html_path)


def process_unsubscribe_requests():
    try:
        imap = imaplib.IMAP4_SSL("imap.mail.ru")
        imap.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        imap.select("INBOX")
        result, data = imap.search(None, '(UNSEEN SUBJECT "unsubscribe")')
        for num in (data[0].split() if data and data[0] else []):
            _, msg_data = imap.fetch(num, "(RFC822)")
            raw_email = msg_data[0][1]
            msg = email.message_from_bytes(raw_email)
            sender = email.utils.parseaddr(msg.get("From"))[1]
            if sender:
                add_blocked_email(sender)
            imap.store(num, '+FLAGS', '\\Seen')
        imap.logout()
    except Exception as e:
        log_error(f"process_unsubscribe_requests: {e}")


# ---------------- Блок-лист ----------------
def get_blocked_emails() -> Set[str]:
    if not os.path.exists(BLOCKED_FILE):
        return set()
    with open(BLOCKED_FILE, "r", encoding="utf-8") as f:
        return set(normalize_email(line) for line in f if "@" in line)


def add_blocked_email(email_str: str) -> bool:
    email_norm = re.sub(r'^\.+', '', normalize_email(email_str))
    if not email_norm or "@" not in email_norm:
        return False
    existing = get_blocked_emails()
    if email_norm in existing:
        return False
    with open(BLOCKED_FILE, "a", encoding="utf-8") as f:
        f.write(email_norm + "\n")
    return True


def dedupe_blocked_file():
    if not os.path.exists(BLOCKED_FILE):
        return
    with open(BLOCKED_FILE, "r", encoding="utf-8") as f:
        raw = [normalize_email(line) for line in f if "@" in line]
    raw = [re.sub(r'^\.+', '', e) for e in raw]
    keep = set(raw)
    by_suffix: Dict[str, Set[str]] = {}
    for e in list(keep):
        m = re.match(r'^(\d{1,2})([A-Za-z][A-Za-z0-9._%+-]*@.+)$', e, flags=re.I)
        if m:
            suffix = m.group(2)
            by_suffix.setdefault(suffix, set()).add(e)
    for suffix, variants in by_suffix.items():
        many = len(variants) >= 2
        clean_present = suffix in keep
        if many or clean_present:
            keep.difference_update(variants)
            keep.add(suffix)
    with open(BLOCKED_FILE, "w", encoding="utf-8") as f:
        if keep:
            f.write("\n".join(sorted(keep)) + "\n")


# ---------------- История отправок ----------------
def log_sent_email(email_addr, group, status="ok", user_id=None, filename=None, error_msg=None):
    os.makedirs(os.path.dirname(LOG_FILE) or ".", exist_ok=True)
    with open(LOG_FILE, "a", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            datetime.now().isoformat(),
            normalize_email(email_addr), group, status,
            user_id if user_id else "",
            filename if filename else "",
            error_msg if error_msg else ""
        ])


def _parse_list_line(line: bytes):
    s = line.decode(errors="ignore")
    m = re.match(r'^\((?P<flags>[^)]*)\)\s+"(?P<delim>[^"]*)"\s+"?(?P<name>.+?)"?$', s)
    if not m:
        return None, ""
    return m.group("name"), m.group("flags")


def detect_sent_folder(imap: imaplib.IMAP4_SSL) -> str:
    status, data = imap.list()
    if status != "OK" or not data:
        return "Sent"
    candidates = []
    for line in data:
        name, flags = _parse_list_line(line)
        if not name:
            continue
        if "\\Sent" in flags or "\\sent" in flags:
            return name
        candidates.append(name)
    prefer = [
        'INBOX.Sent', 'Sent', 'Sent Items', 'Sent Messages',
        'Отправленные', '&BB4EQgQ,BEAEMAQyBDsENQQ9BD0ESwQ1-'
    ]
    for wanted in prefer:
        for c in candidates:
            if c.endswith(wanted) or c == wanted:
                return c
    return candidates[0] if candidates else "Sent"


def get_recent_6m_union() -> Set[str]:
    cutoff = datetime.now() - timedelta(days=180)
    result = set()
    if os.path.exists(LOG_FILE):
        with open(LOG_FILE, encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) < 2:
                    continue
                try:
                    dt = datetime.fromisoformat(row[0])
                    if dt.tzinfo is not None:
                        dt = dt.replace(tzinfo=None)
                except Exception:
                    continue
                if dt >= cutoff:
                    result.add(normalize_email(row[1]))
    return result


def clear_recent_sent_cache():
    _recent_cache["at"] = None
    _recent_cache["set"] = set()


# ---------------- Команды/GUI ----------------
def enable_force_send(chat_id: int) -> None:
    FORCE_SEND_CHAT_IDS.add(chat_id)


def disable_force_send(chat_id: int) -> None:
    FORCE_SEND_CHAT_IDS.discard(chat_id)


def is_force_send(chat_id: int) -> bool:
    return chat_id in FORCE_SEND_CHAT_IDS


def clear_all_awaiting(context: ContextTypes.DEFAULT_TYPE):
    for key in ["awaiting_block_email", "awaiting_manual_email"]:
        context.user_data[key] = False


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        ["📤 Загрузить данные для поиска контактов", "🧹 Очистить список"],
        ["📄 Показать исключения", "🚫 Добавить в исключения"],
        ["✉️ Ручная рассылка", "🧾 О боте"],
        ["🧭 Сменить группу", "📈 Отчёты"],
        ["🔄 Синхронизировать с сервером", "🚀 Игнорировать лимит"]
    ]
    markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text("Можно загрузить данные", reply_markup=markup)


async def prompt_upload(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📥 Загрузите данные с e-mail-адресами для рассылки.\n\n"
        "Поддерживаемые форматы: PDF, Excel (.xlsx), Word (.docx), CSV, ZIP (с этими файлами внутри), а также ссылки на сайты."
    )


async def about_bot(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Бот делает рассылку HTML-писем с учётом истории отправки (IMAP 180 дней) и блок-листа. "
        "Один адрес — не чаще 1 раза в 6 месяцев. Домены: только .ru и .com."
    )


async def add_block_prompt(update: Update, context: ContextTypes.DEFAULT_TYPE):
    clear_all_awaiting(context)
    await update.message.reply_text(
        "Введите email или список email-адресов (через запятую/пробел/с новой строки), которые нужно добавить в исключения:"
    )
    context.user_data["awaiting_block_email"] = True


async def show_blocked_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    dedupe_blocked_file()
    blocked = get_blocked_emails()
    if not blocked:
        await update.message.reply_text("📄 Список исключений пуст.")
    else:
        await update.message.reply_text("📄 В исключениях:\n" + "\n".join(sorted(blocked)))


async def prompt_change_group(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("⚽ Спорт", callback_data="group_спорт")],
        [InlineKeyboardButton("🏕 Туризм", callback_data="group_туризм")],
        [InlineKeyboardButton("🩺 Медицина", callback_data="group_медицина")]
    ]
    await update.message.reply_text("⬇️ Выберите направление рассылки:", reply_markup=InlineKeyboardMarkup(keyboard))


async def force_send_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    enable_force_send(chat_id)
    await update.message.reply_text(
        "Режим игнорирования дневного лимита включён для этого чата.\n"
        "Запустите рассылку ещё раз — ограничение на сегодня будет проигнорировано."
    )


async def report_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("📆 День", callback_data="report_day")],
        [InlineKeyboardButton("🗓 Неделя", callback_data="report_week")],
        [InlineKeyboardButton("🗓 Месяц", callback_data="report_month")],
        [InlineKeyboardButton("📅 Год", callback_data="report_year")],
    ]
    await update.message.reply_text("Выберите период отчёта:", reply_markup=InlineKeyboardMarkup(keyboard))


def get_report(period="day"):
    if not os.path.exists(LOG_FILE):
        return "Нет данных о рассылках."
    now = datetime.now()
    if period == "day":
        start_at = now - timedelta(days=1)
    elif period == "week":
        start_at = now - timedelta(weeks=1)
    elif period == "month":
        start_at = now - timedelta(days=30)
    elif period == "year":
        start_at = now - timedelta(days=365)
    else:
        start_at = now - timedelta(days=1)

    cnt_ok = 0
    cnt_err = 0
    with open(LOG_FILE, encoding="utf-8") as f:
        reader = csv.reader(f)
        for row in reader:
            if len(row) < 4:
                continue
            try:
                dt = datetime.fromisoformat(row[0])
                if dt.tzinfo is not None:
                    dt = dt.replace(tzinfo=None)
            except Exception:
                continue
            if dt >= start_at:
                st = (row[3] or "").strip().lower()
                if st == "ok":
                    cnt_ok += 1
                else:
                    cnt_err += 1
    return f"Успешных: {cnt_ok}\nОшибок: {cnt_err}"


async def report_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    period = query.data.replace("report_", "")
    mapping = {"day": "Отчёт за день", "week": "Отчёт за неделю", "month": "Отчёт за месяц", "year": "Отчёт за год"}
    text = get_report(period)
    await query.edit_message_text(f"📊 {mapping.get(period, period)}:\n{text}")


# ---------------- Синхронизация IMAP ----------------
def get_recently_contacted_emails_cached() -> Set[str]:
    now = time.time()
    at = _recent_cache["at"]
    if at is None or (now - at) > _recent_cache["ttl"]:
        _recent_cache["set"] = get_recent_6m_union()
        _recent_cache["at"] = now
    return _recent_cache["set"]


def sync_log_with_imap() -> int:
    try:
        imap = imaplib.IMAP4_SSL("imap.mail.ru")
        imap.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        sent_folder = detect_sent_folder(imap)
        status, _ = imap.select(f'"{sent_folder}"')
        if status != "OK":
            log_error("Sent folder not selected")
            imap.logout(); return 0
        existing = get_recent_6m_union()
        date_180 = (datetime.now() - timedelta(days=180)).strftime("%d-%b-%Y")
        result, data = imap.search(None, f'SINCE {date_180}')
        added = 0
        for num in (data[0].split() if data and data[0] else []):
            _, msg_data = imap.fetch(num, "(RFC822)")
            raw_email = msg_data[0][1]
            msg = email.message_from_bytes(raw_email)
            to_addr = email.utils.parseaddr(msg.get("To"))[1]
            date_str = msg.get("Date")
            if not to_addr:
                continue
            to_addr = normalize_email(to_addr)
            if to_addr not in existing and is_allowed_tld(to_addr):
                if date_str:
                    try:
                        dt = email.utils.parsedate_to_datetime(date_str)
                        if dt.tzinfo is not None:
                            dt = dt.replace(tzinfo=None)
                        if dt < datetime.now() - timedelta(days=180):
                            continue
                    except Exception:
                        continue
                else:
                    continue
                with open(LOG_FILE, "a", encoding="utf-8", newline="") as f:
                    writer = csv.writer(f)
                    writer.writerow([dt.isoformat(), to_addr, "imap_sync", "external"])
                added += 1
        imap.logout()
        return added
    except Exception as e:
        log_error(f"sync_log_with_imap: {e}")
        raise


async def sync_imap_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("⏳ Сканируем папку «Отправленные» (последние 180 дней)...")
    try:
        added = sync_log_with_imap()
        clear_recent_sent_cache()
        await update.message.reply_text(f"🔄 Добавлено в лог {added} новых адресов.")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка синхронизации: {e}")


# ---------------- Ввод/обработка ----------------
async def reset_email_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    bucket = session_data.setdefault(chat_id, {})
    bucket["all_emails"] = set()
    bucket["all_files"] = []
    bucket["to_send"] = []
    bucket["suspect_numeric"] = []
    bucket["foreign"] = []
    bucket["preview_allowed_all"] = []
    bucket["repairs"] = []
    await update.message.reply_text("Список email-адресов и файлов очищен. Можно загружать новые файлы!")


async def _compose_report_and_save(chat_id: int, allowed_all: Set[str], filtered: List[str],
                                   suspicious_numeric: List[str], foreign: List[str]) -> str:
    bucket = session_data.setdefault(chat_id, {})
    bucket["preview_allowed_all"] = sorted(allowed_all)
    bucket["suspect_numeric"] = suspicious_numeric
    bucket["foreign"] = sorted(foreign)

    sample_allowed = sample_preview(bucket["preview_allowed_all"], PREVIEW_ALLOWED)
    sample_numeric = sample_preview(suspicious_numeric, PREVIEW_NUMERIC)
    sample_foreign = sample_preview(bucket["foreign"], PREVIEW_FOREIGN)

    report = (
        "✅ Анализ завершён.\n"
        f"Найдено адресов (.ru/.com): {len(allowed_all)}\n"
        f"Уникальных (после базовой очистки): {len(filtered)}\n"
        f"Подозрительные (логин только из цифр, исключены): {len(suspicious_numeric)}\n"
        f"Иностранные домены (исключены): {len(foreign)}"
    )
    if sample_allowed:
        report += "\n\n🧪 Примеры (.ru/.com):\n" + "\n".join(sample_allowed)
    if sample_numeric:
        report += "\n\n🔢 Примеры цифровых (исключены):\n" + "\n".join(sample_numeric)
    if sample_foreign:
        report += "\n\n🌍 Примеры иностранных (исключены):\n" + "\n".join(sample_foreign)
    return report


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    if not doc:
        return
    chat_id = update.effective_chat.id
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    file_path = os.path.join(DOWNLOAD_DIR, f"{chat_id}_{int(time.time())}_{doc.file_name}")
    f = await doc.get_file()
    await f.download_to_drive(file_path)

    await update.message.reply_text("Файл загружен. Идёт анализ...")
    progress_msg = await update.message.reply_text("🔎 Анализируем...")

    allowed_all, loose_all = set(), set()
    extracted_files: List[str] = []
    repairs: List[tuple[str, str]] = []

    try:
        if file_path.lower().endswith(".zip"):
            allowed_all, extracted_files, loose_all = await extract_emails_from_zip(file_path, progress_msg, DOWNLOAD_DIR)
            repairs = collect_repairs_from_files(extracted_files)
        else:
            allowed, loose = extract_from_uploaded_file(file_path)
            allowed_all.update(allowed); loose_all.update(loose)
            extracted_files.append(file_path)
            repairs = collect_repairs_from_files([file_path])
    except Exception as e:
        log_error(f"handle_document: {file_path}: {e}")

    # Устраняем «усечённые» цифровые (33@ → vilena33@), ДО отчёта
    allowed_all, trunc_pairs = apply_numeric_truncation_removal(allowed_all)
    repairs = list(dict.fromkeys(repairs + trunc_pairs))

    # Фильтрация технич.
    technical_emails = [e for e in allowed_all if any(tp in e for tp in TECH_PATTERNS)]
    filtered = [e for e in allowed_all if e not in technical_emails and is_allowed_tld(e)]

    # Цифровые логины — исключаем, но показываем (кроме уже сопоставленных trunc_pairs)
    suspicious_numeric = sorted({e for e in filtered if is_numeric_localpart(e)})
    filtered = [e for e in filtered if not is_numeric_localpart(e)]

    # Иностранные домены (для показа) + схлопывание «сносок»
    foreign_raw = {e for e in loose_all if not is_allowed_tld(e)}
    foreign = sorted(collapse_footnote_variants(foreign_raw))

    # Сохраняем в сессию
    bucket = session_data.setdefault(chat_id, {})
    bucket["all_emails"] = set(filtered)
    bucket["all_files"] = extracted_files
    bucket["to_send"] = sorted(set(filtered))
    bucket["group"] = bucket.get("group")
    bucket["template"] = bucket.get("template")
    bucket["repairs"] = repairs
    bucket["repairs_sample"] = sample_preview([f"{b} → {g}" for (b, g) in repairs], 6)

    # Отчёт
    report = await _compose_report_and_save(chat_id, allowed_all, filtered, suspicious_numeric, foreign)
    if repairs:
        report += "\n\n🧩 Возможные исправления (проверьте вручную):"
        for s in bucket["repairs_sample"]:
            report += f"\n{s}"
    try:
        await progress_msg.edit_text(report)
    except Exception:
        pass

    # Доп. действия
    extra_buttons = [[InlineKeyboardButton("🔁 Показать ещё примеры", callback_data="refresh_preview")]]
    if suspicious_numeric:
        extra_buttons.append([InlineKeyboardButton(f"➕ Включить цифровые ({len(suspicious_numeric)})", callback_data="ask_include_numeric")])
        extra_buttons.append([InlineKeyboardButton("🔢 Показать цифровые", callback_data="show_numeric")])
    if foreign:
        extra_buttons.append([InlineKeyboardButton(f"🌍 Показать иностранные ({len(foreign)})", callback_data="show_foreign")])
    if repairs:
        extra_buttons.append([InlineKeyboardButton(f"🧩 Применить исправления ({len(repairs)})", callback_data="apply_repairs")])
        extra_buttons.append([InlineKeyboardButton("🧩 Показать все исправления", callback_data="show_repairs")])
    extra_buttons.append([InlineKeyboardButton("▶️ Перейти к выбору направления", callback_data="proceed_group")])

    await update.message.reply_text(
        "Дополнительные действия:\n"
        "ℹ️ При необходимости: «🔄 Синхронизировать с сервером», затем применить исправления и продолжить.",
        reply_markup=InlineKeyboardMarkup(extra_buttons)
    )


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()

    # режим добавления в исключения
    if context.user_data.get("awaiting_block_email"):
        unique_emails = list(dict.fromkeys(extract_emails_loose(text)))
        if unique_emails:
            added = []
            for e in unique_emails:
                if add_blocked_email(e):
                    added.append(e)
            if added:
                await update.message.reply_text(f"✅ Добавлено в исключения: {', '.join(added)}")
            else:
                await update.message.reply_text("ℹ️ Все указанные адреса уже были в исключениях.")
        else:
            await update.message.reply_text("❌ Не найдено ни одного email.")
        context.user_data["awaiting_block_email"] = False
        return

    # ручная рассылка — ожидаем список адресов
    if context.user_data.get("awaiting_manual_email"):
        emails = list(dict.fromkeys(extract_clean_emails_from_text(text)))
        if emails:
            context.user_data["manual_emails"] = [normalize_email(e) for e in emails]
            context.user_data["awaiting_manual_email"] = False
            keyboard = [
                [InlineKeyboardButton("⚽ Спорт", callback_data="manual_group_спорт")],
                [InlineKeyboardButton("🏕 Туризм", callback_data="manual_group_туризм")],
                [InlineKeyboardButton("🩺 Медицина", callback_data="manual_group_медицина")]
            ]
            await update.message.reply_text(
                f"К отправке: {', '.join(context.user_data['manual_emails'])}\n\n⬇️ Выберите направление письма:",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
        else:
            await update.message.reply_text("❌ Не найдено ни одного email (.ru/.com).")
        return

    # ссылки — парсим
    urls = re.findall(r'https?://\S+', text)
    if urls:
        await update.message.reply_text("📥 Ссылки загружены. Идёт анализ...")
        chat_id = update.effective_chat.id
        async with aiohttp.ClientSession() as session:
            tasks = [async_extract_emails_from_url(url, session, chat_id) for url in urls]
            results = await asyncio.gather(*tasks)

        allowed_all, foreign_all = set(), set()
        repairs_all: List[tuple[str, str]] = []
        for _, allowed, foreign, repairs in results:
            allowed_all.update(allowed); foreign_all.update(foreign); repairs_all.extend(repairs)

        # удаляем усечённые числовые (33@ → name33@)
        allowed_all, trunc_pairs = apply_numeric_truncation_removal(allowed_all)
        repairs_all = list(dict.fromkeys(repairs_all + trunc_pairs))

        technical_emails = [e for e in allowed_all if any(tp in e for tp in TECH_PATTERNS)]
        filtered = [e for e in allowed_all if e not in technical_emails and is_allowed_tld(e)]

        suspicious_numeric = sorted({e for e in filtered if is_numeric_localpart(e)})
        filtered = [e for e in filtered if not is_numeric_localpart(e)]

        foreign_all = collapse_footnote_variants(foreign_all)

        bucket = session_data.setdefault(chat_id, {})
        bucket["all_emails"] = set(filtered)
        bucket["all_files"] = []
        bucket["to_send"] = sorted(set(filtered))
        bucket["group"] = bucket.get("group")
        bucket["template"] = bucket.get("template")
        bucket["repairs"] = repairs_all
        bucket["repairs_sample"] = sample_preview([f"{b} → {g}" for (b, g) in repairs_all], 6)

        report = await _compose_report_and_save(chat_id, allowed_all, filtered, suspicious_numeric, sorted(foreign_all))
        if bucket["repairs"]:
            report += "\n\n🧩 Возможные исправления (проверьте вручную):"
            for s in bucket["repairs_sample"]:
                report += f"\n{s}"
        await update.message.reply_text(report)

        extra_buttons = [[InlineKeyboardButton("🔁 Показать ещё примеры", callback_data="refresh_preview")]]
        if suspicious_numeric:
            extra_buttons.append([InlineKeyboardButton(f"➕ Включить цифровые ({len(suspicious_numeric)})", callback_data="ask_include_numeric")])
            extra_buttons.append([InlineKeyboardButton("🔢 Показать цифровые", callback_data="show_numeric")])
        if bucket["foreign"]:
            extra_buttons.append([InlineKeyboardButton(f"🌍 Показать иностранные ({len(bucket['foreign'])})", callback_data="show_foreign")])
        if bucket["repairs"]:
            extra_buttons.append([InlineKeyboardButton(f"🧩 Применить исправления ({len(bucket['repairs'])})", callback_data="apply_repairs")])
            extra_buttons.append([InlineKeyboardButton("🧩 Показать все исправления", callback_data="show_repairs")])
        extra_buttons.append([InlineKeyboardButton("▶️ Перейти к выбору направления", callback_data="proceed_group")])

        await update.message.reply_text("Дополнительные действия:", reply_markup=InlineKeyboardMarkup(extra_buttons))
        return

    await prompt_upload(update, context)


async def refresh_preview(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    bucket = session_data.get(chat_id, {}) or {}
    allowed_all = bucket.get("preview_allowed_all", [])
    numeric = bucket.get("suspect_numeric", []) or []
    foreign = bucket.get("foreign", []) or []
    if not (allowed_all or numeric or foreign):
        await query.answer("Нет данных для примеров. Загрузите файл/ссылки.", show_alert=True); return
    await query.answer()
    sample_allowed = sample_preview(allowed_all, PREVIEW_ALLOWED)
    sample_numeric = sample_preview(numeric, PREVIEW_NUMERIC)
    sample_foreign = sample_preview(foreign, PREVIEW_FOREIGN)
    report = []
    if sample_allowed: report.append("🧪 Примеры (.ru/.com):\n" + "\n".join(sample_allowed))
    if sample_numeric: report.append("🔢 Примеры цифровых (исключены):\n" + "\n".join(sample_numeric))
    if sample_foreign: report.append("🌍 Примеры иностранных (исключены):\n" + "\n".join(sample_foreign))
    await query.message.reply_text("\n\n".join(report) if report else "Показать нечего.")


async def proceed_to_group(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    keyboard = [
        [InlineKeyboardButton("⚽ Спорт", callback_data="group_спорт")],
        [InlineKeyboardButton("🏕 Туризм", callback_data="group_туризм")],
        [InlineKeyboardButton("🩺 Медицина", callback_data="group_медицина")]
    ]
    await query.message.reply_text("⬇️ Выберите направление рассылки:", reply_markup=InlineKeyboardMarkup(keyboard))


async def select_group(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    group_code = query.data.split("_")[1]
    template_path = TEMPLATE_MAP[group_code]
    chat_id = query.message.chat.id
    emails = session_data.get(chat_id, {}).get("to_send", [])
    session_data.setdefault(chat_id, {})
    session_data[chat_id]["group"] = group_code
    session_data[chat_id]["template"] = template_path
    session_data[chat_id]["to_send"] = emails
    await query.message.reply_text(
        f"✉️ Готово к отправке {len(emails)} писем.\n"
        f"Для запуска рассылки нажмите кнопку ниже.",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("✉️ Начать рассылку", callback_data="start_sending")]])
    )


async def prompt_manual_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    clear_all_awaiting(context)
    await update.message.reply_text("Введите email или список email-адресов (через запятую/пробел/с новой строки):")
    context.user_data["awaiting_manual_email"] = True


async def ask_include_numeric(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    bucket = session_data.setdefault(chat_id, {})
    numeric = bucket.get("suspect_numeric", []) or []
    if not numeric:
        await query.answer("Цифровых адресов нет", show_alert=True); return
    await query.answer()
    preview_list = numeric[:60]
    txt = f"Найдено цифровых логинов: {len(numeric)}.\nБудут добавлены все.\n\nПример:\n" + "\n".join(preview_list)
    more = len(numeric) - len(preview_list)
    if more > 0:
        txt += f"\n… и ещё {more}."
    await query.message.reply_text(
        txt,
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("✅ Включить все цифровые", callback_data="confirm_include_numeric")],
            [InlineKeyboardButton("↩️ Отмена", callback_data="cancel_include_numeric")]
        ])
    )


async def include_numeric_emails(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    bucket = session_data.setdefault(chat_id, {})
    numeric = bucket.get("suspect_numeric", []) or []
    if not numeric:
        await query.answer("Цифровых адресов нет", show_alert=True); return
    current = set(bucket.get("to_send", []))
    added = [e for e in numeric if e not in current]
    current.update(numeric)
    bucket["to_send"] = sorted(current)
    await query.answer()
    await query.message.reply_text(f"➕ Добавлено цифровых адресов: {len(added)}.\nИтого к отправке: {len(bucket['to_send'])}.")


async def cancel_include_numeric(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text("Ок, цифровые адреса оставлены выключенными.")


def _chunk_list(items: List[str], size=60) -> List[List[str]]:
    return [items[i:i+size] for i in range(0, len(items), size)]


async def show_numeric_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    numeric = session_data.get(chat_id, {}).get("suspect_numeric", []) or []
    if not numeric:
        await query.answer("Список пуст", show_alert=True); return
    await query.answer()
    for chunk in _chunk_list(numeric, 60):
        await query.message.reply_text("🔢 Цифровые логины:\n" + "\n".join(chunk))


async def show_foreign_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    foreign = session_data.get(chat_id, {}).get("foreign", []) or []
    if not foreign:
        await query.answer("Список пуст", show_alert=True); return
    await query.answer()
    for chunk in _chunk_list(foreign, 60):
        await query.message.reply_text("🌍 Иностранные домены (исключены):\n" + "\n".join(chunk))


async def apply_repairs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    bucket = session_data.setdefault(chat_id, {})
    repairs: List[tuple[str, str]] = bucket.get("repairs", []) or []
    if not repairs:
        await query.answer("Нет кандидатов на исправление", show_alert=True); return
    current = set(bucket.get("to_send", []))
    applied = 0
    changed = []
    for bad, good in repairs:
        if bad in current:
            current.discard(bad)
            if is_allowed_tld(good):
                current.add(good)
                applied += 1
                if applied <= 12:
                    changed.append(f"{bad} → {good}")
    bucket["to_send"] = sorted(current)
    await query.answer()
    txt = f"🧩 Применено исправлений: {applied}."
    if changed:
        txt += "\n" + "\n".join(changed)
        if applied > len(changed):
            txt += f"\n… и ещё {applied - len(changed)}."
    await query.message.reply_text(txt)


async def show_repairs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    repairs: List[tuple[str, str]] = session_data.get(chat_id, {}).get("repairs", []) or []
    if not repairs:
        await query.answer("Список пуст", show_alert=True); return
    await query.answer()
    pairs = [f"{b} → {g}" for (b, g) in repairs]
    for chunk in _chunk_list(pairs, 60):
        await query.message.reply_text("🧩 Возможные исправления:\n" + "\n".join(chunk))


async def send_manual_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    group_code = query.data.split("_")[2]
    template_path = TEMPLATE_MAP[group_code]

    emails_raw = context.user_data.get("manual_emails", [])
    all_text = " ".join(emails_raw)
    emails = sorted({normalize_email(x) for x in extract_clean_emails_from_text(all_text)})
    if not emails:
        await query.message.reply_text("❗ Список email пуст."); return

    loop = asyncio.get_running_loop()
    recent_sent = await loop.run_in_executor(None, get_recent_6m_union)
    blocked = get_blocked_emails()
    sent_today = get_sent_today()

    to_send = [e for e in emails if e not in recent_sent and e not in sent_today and e not in blocked]
    if not to_send:
        await query.message.reply_text("❗ Все адреса уже есть в истории за 6 мес. или в блок-листе.")
        context.user_data["manual_emails"] = []; return

    available = max(0, MAX_EMAILS_PER_DAY - len(sent_today))
    if available <= 0 and not is_force_send(chat_id):
        await update.callback_query.message.reply_text(
            f"❗ Дневной лимит {MAX_EMAILS_PER_DAY} уже исчерпан.\n"
            "Если вы исправили ошибки — нажмите «🚀 Игнорировать лимит» и запустите ещё раз."
        ); return
    if not is_force_send(chat_id) and len(to_send) > available:
        to_send = to_send[:available]
        await query.message.reply_text(f"⚠️ Учитываю дневной лимит: будет отправлено {available} адресов из списка.")

    await query.message.reply_text(f"✉️ Рассылка начата. Отправляем {len(to_send)} писем...")

    sent_count = 0
    errors = []
    for email_addr in to_send:
        try:
            await async_send_email(email_addr, template_path)
            log_sent_email(email_addr, group_code, "ok", chat_id, template_path)
            sent_count += 1
            await asyncio.sleep(1.5)
        except Exception as e:
            errors.append(f"{email_addr} — {e}")
            err = str(e).lower()
            if ("invalid mailbox" in err or "user is terminated" in err or "non-local recipient verification failed" in err):
                add_blocked_email(email_addr)
            log_sent_email(email_addr, group_code, "error", chat_id, template_path, str(e))

    await query.message.reply_text(f"✅ Отправлено писем: {sent_count}")
    if errors:
        await query.message.reply_text("Ошибки:\n" + "\n".join(errors))

    context.user_data["manual_emails"] = []
    clear_recent_sent_cache()
    disable_force_send(chat_id)


def get_sent_today() -> Set[str]:
    if not os.path.exists(LOG_FILE):
        return set()
    today = datetime.now().date()
    sent_today = set()
    with open(LOG_FILE, encoding="utf-8") as f:
        reader = csv.reader(f)
        for row in reader:
            if len(row) < 4:
                continue
            try:
                dt = datetime.fromisoformat(row[0])
                if dt.tzinfo is not None:
                    dt = dt.replace(tzinfo=None)
            except Exception:
                continue
            if dt.date() != today:
                continue
            status = (row[3] or "").strip().lower()
            if status == "ok":
                sent_today.add(normalize_email(row[1]))
    return sent_today


async def send_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = query.message.chat.id
    data = session_data.get(chat_id, {})
    to_send_raw = data.get("to_send", [])
    template_path = data.get("template")
    group_code = data.get("group")

    all_text = " ".join(to_send_raw)
    to_send = sorted({normalize_email(x) for x in extract_clean_emails_from_text(all_text)})

    if not to_send:
        await query.message.reply_text("❗ Нет email для рассылки."); return

    loop = asyncio.get_running_loop()
    recent_sent = await loop.run_in_executor(None, get_recent_6m_union)
    blocked = get_blocked_emails()
    sent_today = get_sent_today()

    emails_to_send = [e for e in to_send if e not in recent_sent and e not in sent_today and e not in blocked]
    if not emails_to_send:
        await query.message.reply_text("❗ Все адреса уже есть в истории за 6 мес. или в блок-листе."); return

    available = max(0, MAX_EMAILS_PER_DAY - len(sent_today))
    if available <= 0 and not is_force_send(chat_id):
        await query.message.reply_text(
            f"❗ Дневной лимит {MAX_EMAILS_PER_DAY} уже исчерпан.\n"
            "Если вы исправили ошибки — нажмите «🚀 Игнорировать лимит» и запустите ещё раз."
        ); return
    if not is_force_send(chat_id) and len(emails_to_send) > available:
        emails_to_send = emails_to_send[:available]
        await query.message.reply_text(f"⚠️ Учитываю дневной лимит: будет отправлено {available} адресов из списка.")

    await query.message.reply_text(f"✉️ Рассылка начата. Отправляем {len(emails_to_send)} писем...")

    sent_count = 0
    bad_emails = []
    for email_addr in emails_to_send:
        try:
            await async_send_email(email_addr, template_path)
            log_sent_email(email_addr, group_code, "ok", chat_id, template_path)
            sent_count += 1
            await asyncio.sleep(1.5)
        except Exception as e:
            error_text = str(e).lower()
            if ("invalid mailbox" in error_text or "user is terminated" in error_text or "non-local recipient verification failed" in error_text):
                add_blocked_email(email_addr)
                bad_emails.append(email_addr)
            log_sent_email(email_addr, group_code, "error", chat_id, template_path, str(e))

    await query.message.reply_text(f"✅ Отправлено писем: {sent_count}")
    if bad_emails:
        await query.message.reply_text("🚫 В блок-лист добавлены:\n" + "\n".join(bad_emails))

    clear_recent_sent_cache()
    disable_force_send(chat_id)


# ---------------- Сервис/потоки ----------------
async def autosync_imap_with_message(query):
    await query.message.reply_text("🔄 Синхронизация истории отправки с сервером...")
    loop = asyncio.get_running_loop()
    added = await loop.run_in_executor(None, sync_log_with_imap)
    clear_recent_sent_cache()
    await query.message.reply_text(
        f"✅ Синхронизация завершена. В лог добавлено новых адресов: {added}.\n"
        f"История отправки обновлена на последние 6 месяцев."
    )


def periodic_unsubscribe_check():
    while True:
        try:
            process_unsubscribe_requests()
        except Exception as e:
            log_error(f"periodic_unsubscribe_check: {e}")
        time.sleep(300)


def check_env_vars():
    for var in ["TELEGRAM_BOT_TOKEN", "EMAIL_ADDRESS", "EMAIL_PASSWORD"]:
        if not os.getenv(var):
            raise EnvironmentError(f"Переменная окружения {var} не задана!")


# ---------------- Запуск ----------------
def main():
    load_env()
    global TOKEN, EMAIL_ADDRESS, EMAIL_PASSWORD
    TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
    EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS", "")
    EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD", "")
    check_env_vars()

    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    dedupe_blocked_file()

    app = ApplicationBuilder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", start))

    # Главное меню
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^📤"), prompt_upload))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^🧹"), reset_email_list))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^🧾"), about_bot))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^🚫"), add_block_prompt))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^📄"), show_blocked_list))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^✉️"), prompt_manual_email))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^🧭"), prompt_change_group))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^📈"), report_command))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^🔄"), sync_imap_command))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex("^🚀"), force_send_command))

    # Файлы и тексты
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))  # последним!

    # Колбэки
    app.add_handler(CallbackQueryHandler(send_manual_email, pattern="^manual_group_"))
    app.add_handler(CallbackQueryHandler(proceed_to_group, pattern="^proceed_group$"))
    app.add_handler(CallbackQueryHandler(select_group, pattern="^group_"))
    app.add_handler(CallbackQueryHandler(send_all, pattern="^start_sending"))
    app.add_handler(CallbackQueryHandler(report_callback, pattern="^report_"))
    app.add_handler(CallbackQueryHandler(show_numeric_list, pattern="^show_numeric$"))
    app.add_handler(CallbackQueryHandler(show_foreign_list, pattern="^show_foreign$"))
    app.add_handler(CallbackQueryHandler(refresh_preview, pattern="^refresh_preview$"))
    app.add_handler(CallbackQueryHandler(ask_include_numeric, pattern="^ask_include_numeric$"))
    app.add_handler(CallbackQueryHandler(include_numeric_emails, pattern="^confirm_include_numeric$"))
    app.add_handler(CallbackQueryHandler(cancel_include_numeric, pattern="^cancel_include_numeric$"))
    app.add_handler(CallbackQueryHandler(apply_repairs, pattern="^apply_repairs$"))
    app.add_handler(CallbackQueryHandler(show_repairs,  pattern="^show_repairs$"))

    print("Бот запущен.")
    t = threading.Thread(target=periodic_unsubscribe_check, daemon=True)
    t.start()
    app.run_polling()


if __name__ == "__main__":
    main()
